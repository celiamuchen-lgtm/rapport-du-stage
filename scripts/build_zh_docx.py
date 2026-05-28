from pathlib import Path
import re
import argparse

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SRC = ROOT / "03_manuscript" / "zh_working" / "09_rapport_stage_premier_jet.zh.md"
DEFAULT_OUT = ROOT / "06_outputs" / "09_rapport_stage_premier_jet.zh.docx"

BODY_FONT = "SimSun"
HEAD_FONT = "Microsoft YaHei"
LATIN_FONT = "Times New Roman"
ACCENT = RGBColor(31, 95, 107)
MUTED = RGBColor(95, 95, 95)
TODO = RGBColor(163, 58, 58)
LIGHT = "EAF3F5"


def set_run_font(run, east_asia=BODY_FONT, latin=LATIN_FONT, size=None, bold=False, color=None):
    run.font.name = latin
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, size=10, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def clean_inline(text):
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    return text.strip()


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "右键更新域以刷新目录"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    set_run_font(r, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r._r.extend([fld_begin, instr, fld_end])


def ensure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(6)

    for i, size in [(1, 18), (2, 15), (3, 12)]:
        style = styles[f"Heading {i}"]
        style.font.name = LATIN_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT if i < 3 else RGBColor(40, 40, 40)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEAD_FONT)
        style.paragraph_format.space_before = Pt(14 if i == 1 else 9)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    if "TODO Box" not in styles:
        todo = styles.add_style("TODO Box", WD_STYLE_TYPE.PARAGRAPH)
        todo.font.name = LATIN_FONT
        todo.font.size = Pt(10.5)
        todo.font.bold = True
        todo.font.color.rgb = TODO
        todo._element.rPr.rFonts.set(qn("w:eastAsia"), HEAD_FONT)
        todo.paragraph_format.left_indent = Cm(0.4)
        todo.paragraph_format.right_indent = Cm(0.2)
        todo.paragraph_format.space_before = Pt(5)
        todo.paragraph_format.space_after = Pt(8)


def add_paragraph(doc, text, style=None, indent=True):
    p = doc.add_paragraph(style=style)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(clean_inline(text))
    is_todo = text.startswith("TODO")
    set_run_font(run, size=10.8, bold=is_todo, color=TODO if is_todo else None)
    return p


def add_cover(doc, lines):
    title = clean_inline(lines[0][2:].strip())
    subtitle = clean_inline(lines[1][3:].strip()) if len(lines) > 1 and lines[1].startswith("##") else ""
    metadata = [clean_inline(x) for x in lines[2:] if x.strip() and not x.startswith("**")]

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, east_asia=HEAD_FONT, latin=LATIN_FONT, size=23, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run_font(r, east_asia=HEAD_FONT, latin=LATIN_FONT, size=15, bold=True, color=RGBColor(60, 60, 60))

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("中文修改版")
    set_run_font(r, east_asia=HEAD_FONT, latin=LATIN_FONT, size=13, bold=True)

    for item in metadata:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(item)
        set_run_font(r, size=11, color=MUTED)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目录")
    set_run_font(r, east_asia=HEAD_FONT, size=18, bold=True, color=ACCENT)
    p = doc.add_paragraph()
    add_field(p, r'TOC \o "1-3" \h \z \u')
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def parse_table(lines):
    rows = []
    for line in lines:
        cells = [clean_inline(x.strip()) for x in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", x.replace(" ", "")) for x in rows[1]):
        rows.pop(1)
    return rows


def add_table(doc, table_lines):
    rows = parse_table(table_lines)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            set_cell_text(cell, val, bold=i == 0)
            if i == 0:
                set_cell_shading(cell, LIGHT)
    doc.add_paragraph()


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--src", type=Path, default=DEFAULT_SRC)
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT)
    args = parser.parse_args()
    src = args.src if args.src.is_absolute() else ROOT / args.src
    out = args.out if args.out.is_absolute() else ROOT / args.out

    out.parent.mkdir(exist_ok=True)
    lines = src.read_text(encoding="utf-8").splitlines()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.4)
    add_page_number(section)
    ensure_styles(doc)

    cover_lines = []
    idx = 0
    while idx < len(lines) and not (idx > 0 and lines[idx].startswith("# ")):
        if lines[idx].strip():
            cover_lines.append(lines[idx])
        idx += 1
    add_cover(doc, cover_lines)

    table_buffer = []
    first_h1 = True

    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            add_table(doc, table_buffer)
            table_buffer = []

    while idx < len(lines):
        raw = lines[idx].rstrip()
        idx += 1
        if not raw.strip():
            flush_table()
            continue

        if raw.startswith("|"):
            table_buffer.append(raw)
            continue
        flush_table()

        if raw.startswith("#"):
            level = len(raw) - len(raw.lstrip("#"))
            text = clean_inline(raw[level:].strip())
            if level == 1 and not first_h1:
                doc.add_section(WD_SECTION.NEW_PAGE)
            first_h1 = False
            p = doc.add_heading(text, level=min(level, 3))
            for run in p.runs:
                set_run_font(run, east_asia=HEAD_FONT, size={1: 18, 2: 15, 3: 12}.get(level, 11), bold=True, color=ACCENT if level < 3 else RGBColor(45, 45, 45))
            continue

        if raw.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.7)
            r = p.add_run(clean_inline(raw[2:]))
            set_run_font(r, size=10.8)
            continue

        if re.match(r"^\d+\.\s", raw):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(0.7)
            r = p.add_run(clean_inline(re.sub(r"^\d+\.\s", "", raw)))
            set_run_font(r, size=10.8)
            continue

        if raw.startswith("TODO"):
            p = add_paragraph(doc, raw, style="TODO Box", indent=False)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        add_paragraph(doc, raw)

    flush_table()
    doc.save(out)
    print(f"Generated: {out}")


if __name__ == "__main__":
    main()
