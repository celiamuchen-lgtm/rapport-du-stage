from pathlib import Path
import re

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "03_manuscript"
OUTPUTS = ROOT / "06_outputs"

CHAPTERS = [
    "00_front_matter.md",
    "01_introduction.md",
    "02_partie_1_structure_territoire.md",
    "03_partie_2_missions.md",
    "04_partie_3_analyse_innovation.md",
    "05_partie_4_bilan.md",
    "06_conclusion.md",
    "07_bibliographie.md",
]


def iter_blocks(markdown: str):
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if not line:
            yield ("blank", "")
        elif line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            yield ("heading", (level, text))
        elif line.startswith("- "):
            yield ("bullet", line[2:].strip())
        else:
            yield ("paragraph", line)


def clean_inline(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    return text


def add_markdown_to_doc(doc: Document, markdown: str):
    for kind, value in iter_blocks(markdown):
        if kind == "blank":
            continue
        if kind == "heading":
            level, text = value
            doc.add_heading(clean_inline(text), level=min(level, 3))
        elif kind == "bullet":
            doc.add_paragraph(clean_inline(value), style="List Bullet")
        else:
            doc.add_paragraph(clean_inline(value))


def main():
    OUTPUTS.mkdir(exist_ok=True)

    pieces = []
    for name in CHAPTERS:
        path = MANUSCRIPT / name
        pieces.append(path.read_text(encoding="utf-8").strip())

    full = "\n\n".join(pieces) + "\n"
    (MANUSCRIPT / "report_full.md").write_text(full, encoding="utf-8")

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    add_markdown_to_doc(doc, full)

    out = OUTPUTS / "rapport_stage_antibes_draft.docx"
    doc.save(out)
    print(f"Generated: {out}")


if __name__ == "__main__":
    main()
