from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "03_manuscript" / "fr_final" / "12_rapport_stage_premier_jet_presentation.fr.tex"
OUT = ROOT / "06_outputs" / "12_rapport_stage_premier_jet_presentation_editable.fr.docx"


ACCENT = RGBColor(0x1F, 0x5F, 0x6B)


def clean_latex(text: str) -> str:
    text = text.strip()
    replacements = {
        r"\%": "%",
        r"\&": "&",
        r"\_": "_",
        r"\#": "#",
        r"\ldots": "...",
        r"--": "–",
        r"``": "« ",
        r"''": " »",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    text = re.sub(r"\\enquote\{([^{}]*)\}", r"« \1 »", text)
    text = re.sub(r"\\textit\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\emph\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\textbf\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\source\{([^{}]*)\}", r"Sources mobilisées : \1", text)
    text = re.sub(r"\\url\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?(?:\{[^{}]*\})?", "", text)
    text = text.replace("{", "").replace("}", "")
    return re.sub(r"\s+", " ", text).strip()


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Cliquez avec le bouton droit puis mettez à jour le champ."
    fld_text = OxmlElement("w:r")
    fld_text.append(text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    paragraph._p.append(fld_text)
    run._r.append(fld_end)


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.paragraph_format.first_line_indent = Cm(0.7)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.25

    for name, size in [("Heading 1", 18), ("Heading 2", 15), ("Heading 3", 13)]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT if name in {"Heading 1", "Heading 2"} else RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(16)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True


def add_title_page(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("De l’analyse socio-thermique à l’épreuve du terrain")
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Comment l’ingénierie territoriale compose-t-elle avec l’urgence climatique "
        "et les contraintes d’un centre ancien ?"
    )
    run.font.size = Pt(14)

    for text in [
        "",
        "Rapport de stage – première version de présentation",
        "",
        "Ke Muchen",
        "Master Ville et Environnement, parcours GEOPRAD",
        "Université Côte d’Azur",
        "",
        "Stage au sein de la Direction de l’Urbanisme",
        "Mairie d’Antibes Juan-les-Pins",
        "",
        "Stage du 1er avril 2026 au 1er juillet 2026",
    ]:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()


def iter_body_lines() -> list[str]:
    text = SRC.read_text(encoding="utf-8")
    text = text.split(r"\begin{document}", 1)[1].split(r"\end{document}", 1)[0]
    text = re.sub(r"\\begin\{titlepage\}.*?\\end\{titlepage\}", "", text, flags=re.S)
    lines = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line or line.startswith("%"):
            lines.append("")
            continue
        if line in {r"\tableofcontents", r"\clearpage", r"\newpage"}:
            continue
        if line.startswith(r"\addcontentsline"):
            continue
        lines.append(line)
    return lines


def add_body(doc: Document) -> None:
    doc.add_heading("Table des matières", level=1)
    add_toc(doc.add_paragraph())
    doc.add_page_break()

    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        paragraph = clean_latex(" ".join(paragraph_buffer))
        paragraph_buffer.clear()
        if paragraph:
            p = doc.add_paragraph(paragraph)
            if paragraph.startswith("Sources mobilisées :"):
                p.style = doc.styles["Intense Quote"]

    for line in iter_body_lines():
        if not line:
            flush_paragraph()
            continue

        heading = re.match(r"\\chapter\*?\{(.+)\}", line)
        if heading:
            flush_paragraph()
            doc.add_heading(clean_latex(heading.group(1)), level=1)
            continue

        heading = re.match(r"\\section\*?\{(.+)\}", line)
        if heading:
            flush_paragraph()
            doc.add_heading(clean_latex(heading.group(1)), level=2)
            continue

        heading = re.match(r"\\subsection\*?\{(.+)\}", line)
        if heading:
            flush_paragraph()
            doc.add_heading(clean_latex(heading.group(1)), level=3)
            continue

        if line.startswith(r"\begin{itemize}") or line.startswith(r"\end{itemize}"):
            flush_paragraph()
            continue
        if line.startswith(r"\item"):
            flush_paragraph()
            doc.add_paragraph(clean_latex(line.removeprefix(r"\item")), style="List Bullet")
            continue

        paragraph_buffer.append(line)

    flush_paragraph()


def main() -> None:
    doc = Document()
    set_doc_defaults(doc)
    add_title_page(doc)
    add_body(doc)

    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
